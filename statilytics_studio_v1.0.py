

import sys, os, io, re, base64, traceback, textwrap, csv, webbrowser, webbrowser
from typing import List, Optional, Dict, Any

import numpy as np
import html
import pandas as pd

from PySide6 import QtCore, QtGui, QtWidgets
from PySide6.QtCore import Qt
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QFileDialog, QMessageBox, QTableView,
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QGroupBox, QFormLayout,
    QComboBox, QSpinBox, QLineEdit, QTextEdit, QTabWidget, QCheckBox, QAbstractItemView,
    QScrollArea, QSizePolicy, QTextBrowser, QListWidget, QListWidgetItem
)

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import importlib.metadata as ilm

from scipy import stats
import statsmodels.api as sm
import statsmodels.formula.api as smf
from statsmodels.stats.multicomp import pairwise_tukeyhsd
from statsmodels.stats.anova import AnovaRM
from statsmodels.multivariate.manova import MANOVA
from statsmodels.stats.inter_rater import fleiss_kappa as sm_fleiss_kappa

from sklearn.decomposition import PCA
from sklearn.cluster import KMeans, AgglomerativeClustering
from sklearn.discriminant_analysis import LinearDiscriminantAnalysis as LDA
from sklearn.metrics import roc_curve, precision_recall_curve, cohen_kappa_score, silhouette_score
from sklearn.linear_model import LogisticRegression
from sklearn.preprocessing import StandardScaler, LabelEncoder
from sklearn.experimental import enable_iterative_imputer  
from sklearn.impute import IterativeImputer


HAS_LIFELINES = True
try:
    from lifelines import KaplanMeierFitter, CoxPHFitter, WeibullAFTFitter
except Exception:
    HAS_LIFELINES = False

HAS_PINGOUIN = True
try:
    import pingouin as pg
except Exception:
    HAS_PINGOUIN = False

HAS_FA = True
try:
    from factor_analyzer import FactorAnalyzer
except Exception:
    HAS_FA = False

HAS_PRINCE = True
try:
    import prince
except Exception:
    HAS_PRINCE = False

HAS_KRIPP = True
try:
    import krippendorff as kd
except Exception:
    HAS_KRIPP = False

APP_NAME = "Statilytics Studio v1.0"
def _app_dir():
    import sys, os
    try:
        base = os.path.dirname(os.path.abspath(__file__))
    except Exception:
        base = os.path.dirname(os.path.abspath(sys.argv[0]))
    if getattr(sys, "frozen", False):
        base = getattr(sys, "_MEIPASS", base)
    return base

COPYRIGHT = "© 2025 Mirza Niaz Zaman Elin. All rights reserved."
WINDOW_TITLE = f"{APP_NAME} — {COPYRIGHT}"


class PandasModel(QtCore.QAbstractTableModel):
    def __init__(self, df=pd.DataFrame(), parent=None):
        super().__init__(parent); self._df = df
    def rowCount(self, parent=None): return len(self._df.index)
    def columnCount(self, parent=None): return self._df.columns.size
    def data(self, index, role=Qt.DisplayRole):
        if not index.isValid(): return None
        if role == Qt.DisplayRole:
            v = self._df.iloc[index.row(), index.column()]
            return "" if pd.isna(v) else str(v)
        return None
    def headerData(self, section, orientation, role=Qt.DisplayRole):
        if role != Qt.DisplayRole: return None
        return str(self._df.columns[section]) if orientation==Qt.Horizontal else str(self._df.index[section])
    def setDataFrame(self, df):
        self.beginResetModel(); self._df = df.copy(); self.endResetModel()


class ReportBuilder:
    def __init__(self):
        self.parts = []
        self.parts.append(textwrap.dedent(f"""
        <html><head><meta charset='utf-8'>
        <style>
        body {{ font-family: -apple-system, Segoe UI, Roboto, Helvetica, Arial, sans-serif; margin: 22px; }}
        .card {{ border: 1px solid #e5e7eb; border-radius: 12px; padding: 14px; margin: 12px 0; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th,td {{ border: 1px solid #e5e7eb; padding: 6px; text-align: left; white-space: nowrap; }}
        th {{ background: #f9fafb; }}
        img {{ max-width: 100%; height: auto; }}
        .muted {{ color:#6b7280; }}
        </style></head><body>
        <h1>{APP_NAME}</h1>
        <p class='muted'>{COPYRIGHT}</p>
        """))
    def add_info(self, title, body): self.parts.append(f"<div class='card'><h3>{title}</h3><div>{body}</div></div>")
    def add_kv(self, title, kv: Dict[str,Any]):
        rows = "".join([f"<tr><th style='width:280px'>{k}</th><td>{v}</td></tr>" for k,v in kv.items()])
        self.parts.append(f"<div class='card'><h3>{title}</h3><table>{rows}</table></div>")
    def add_table(self, df: pd.DataFrame, title="Table"):
        if df is None or len(df)==0: self.add_info(title, "No rows."); return
        df = df.copy()
        try:
            df = df.replace({np.nan:""})
        except Exception:
            pass
        thead = "".join([f"<th>{str(c)}</th>" for c in df.columns])
        rows = []
        for _,r in df.iterrows(): rows.append("<tr>"+"".join([f"<td>{str(v)}</td>" for v in r])+ "</tr>")
        self.parts.append(f"<div class='card'><h3>{title}</h3><table><thead><tr>{thead}</tr></thead><tbody>{''.join(rows)}</tbody></table></div>")
    def add_figure(self, fig, title="Figure"):
        buf = io.BytesIO(); fig.savefig(buf, format='png', dpi=160, bbox_inches='tight'); plt.close(fig)
        b64 = base64.b64encode(buf.getvalue()).decode('ascii')
        self.parts.append(f"<div class='card'><h3>{title}</h3><img src='data:image/png;base64,{b64}'/></div>")
    def html(self): return "\n".join(self.parts + [f"<hr><p class='muted'>{APP_NAME} — {COPYRIGHT}</p></body></html>"])


class Engine:
    def __init__(self, df: pd.DataFrame): self.df = df.copy()

    
    def describe(self): d = self.df.describe(include='all').T; d['missing'] = self.df.isna().sum(); return d
    def crosstab(self, a: str, b: str, fisher=False):
        tab = pd.crosstab(self.df[a], self.df[b])
        chi2, p, dof, _ = stats.chi2_contingency(tab, correction=False)
        info = {"Chi-square": round(chi2,4), "p": round(p,6), "df": dof}
        if fisher and tab.shape==(2,2): _, pf = stats.fisher_exact(tab); info["Fisher p"] = round(pf,6)
        return tab, info

    
    def t_one_sample(self, col: str, mu=0.0):
        x = self.df[col].dropna().astype(float); t, p = stats.ttest_1samp(x, mu)
        return {"n": len(x), "mean": x.mean(), "sd": x.std(ddof=1), "t": t, "p": p}
    def t_ind(self, col: str, group: str, g1, g2, equal_var=False):
        x1 = self.df.loc[self.df[group]==g1, col].dropna().astype(float)
        x2 = self.df.loc[self.df[group]==g2, col].dropna().astype(float)
        t,p = stats.ttest_ind(x1, x2, equal_var=equal_var); return {"n1": len(x1), "n2": len(x2), "t": t, "p": p}
    def t_paired(self, col1: str, col2: str):
        x1 = self.df[col1].astype(float); x2 = self.df[col2].astype(float)
        m = ~(x1.isna() | x2.isna()); x1=x1[m]; x2=x2[m]; t,p = stats.ttest_rel(x1, x2); return {"n": len(x1), "t": t, "p": p}
    def anova_oneway(self, dv: str, group: str):
        groups = [g.dropna().astype(float).values for _,g in self.df[[dv,group]].dropna().groupby(group)[dv]]
        F,p = stats.f_oneway(*groups); return F,p
    def tukey_hsd(self, dv: str, group: str):
        d = self.df[[dv,group]].dropna(); res = pairwise_tukeyhsd(endog=d[dv].astype(float), groups=d[group].astype(str))
        return pd.DataFrame(res._results_table.data[1:], columns=res._results_table.data[0])
    def ancova(self, dv: str, factor: str, covars: List[str]):
        data = self.df.dropna(subset=[dv,factor]+covars); m = smf.ols(f"{dv} ~ C({factor}) + " + " + ".join(covars), data=data).fit(); return m
    def mannwhitney(self, dv: str, group: str, g1, g2):
        x1 = self.df.loc[self.df[group]==g1, dv].dropna().astype(float)
        x2 = self.df.loc[self.df[group]==g2, dv].dropna().astype(float)
        u,p = stats.mannwhitneyu(x1, x2); return {"U": u, "p": p}
    def wilcoxon(self, col1: str, col2: str):
        x1 = self.df[col1].astype(float); x2 = self.df[col2].astype(float)
        m = ~(x1.isna() | x2.isna()); x1=x1[m]; x2=x2[m]; W,p = stats.wilcoxon(x1, x2); return {"W": W, "p": p}
    def kruskal(self, dv: str, group: str):
        groups = [g.dropna().astype(float).values for _,g in self.df[[dv,group]].dropna().groupby(group)[dv]]
        H,p = stats.kruskal(*groups); return {"H": H, "p": p}
    def friedman(self, cols: List[str]):
        arrs = [self.df[c].astype(float).dropna().values for c in cols]; m = min(map(len, arrs)); arrs = [a[:m] for a in arrs]
        Q,p = stats.friedmanchisquare(*arrs); return {"Q": Q, "p": p}

    
    def correlations(self, cols: List[str], method='pearson'): return self.df[cols].corr(method=method)

    
    def ols(self, dv: str, predictors: List[str]):
        data = self.df.dropna(subset=[dv]+predictors); return smf.ols(f"{dv} ~ " + " + ".join(predictors), data=data).fit()
    def logit(self, dv: str, predictors: List[str]):
        data = self.df.dropna(subset=[dv]+predictors); return smf.logit(f"{dv} ~ " + " + ".join(predictors), data=data).fit(disp=False)
    def mlogit(self, dv: str, predictors: List[str]):
        data = self.df.dropna(subset=[dv]+predictors); return smf.mnlogit(f"{dv} ~ " + " + ".join(predictors), data=data).fit(method='newton', maxiter=100, disp=False)
    def ordinal(self, dv: str, predictors: List[str]):
        from statsmodels.miscmodels.ordinal_model import OrderedModel
        data = self.df.dropna(subset=[dv]+predictors); return OrderedModel(data[dv], sm.add_constant(data[predictors]), distr='logit').fit(method='bfgs', disp=False)

    def glm(self, dv: str, predictors: List[str], family: str):
        fam = {'poisson': sm.families.Poisson(), 'negbin': sm.families.NegativeBinomial(), 'gamma': sm.families.Gamma(link=sm.genmod.families.links.log())}[family]
        data = self.df.dropna(subset=[dv]+predictors); return smf.glm(f"{dv} ~ " + " + ".join(predictors), data=data, family=fam).fit()

    
    def lmm(self, dv: str, predictors: List[str], group: str):
        data = self.df.dropna(subset=[dv,group]+predictors); return smf.mixedlm(f"{dv} ~ " + " + ".join(predictors), data=data, groups=data[group]).fit()

    
    def anova_rm(self, dv: str, subject: str, within: str):
        data = self.df.dropna(subset=[dv, subject, within]).copy(); res = AnovaRM(data, depvar=dv, subject=subject, within=[within]).fit(); return res.anova_table
    def manova(self, y_cols: List[str], predictors: List[str]):
        data = self.df.dropna(subset=y_cols+predictors).copy(); y = " + ".join(y_cols); x = " + ".join(predictors) if predictors else "1"
        return MANOVA.from_formula(f"{y} ~ {x}", data=data).mv_test()

    
    def cronbach_alpha(self, cols: List[str]):
        X = self.df[cols].dropna().astype(float).values
        k = X.shape[1]; var_sum = X.var(axis=0, ddof=1).sum(); total_var = X.sum(axis=1).var(ddof=1)
        return (k/(k-1.0))*(1 - var_sum/total_var)

    def kappa(self, col1: str, col2: str):
        if not col1 or not col2:
            raise ValueError("Select both Y and X columns for Cohen's kappa.")
        for c in (col1, col2):
            if c not in self.df.columns:
                raise ValueError(f"Column '{c}' not in dataset. Available: {list(self.df.columns)}")
        a = self.df[col1].astype(str); b = self.df[col2].astype(str)
        n = min(len(a), len(b)); a=a.iloc[:n]; b=b.iloc[:n]
        return cohen_kappa_score(a, b)

    def weighted_kappa(self, col1: str, col2: str, weights: str = "linear"):
        if not col1 or not col2:
            raise ValueError("Select both Y and X columns for weighted kappa.")
        for c in (col1, col2):
            if c not in self.df.columns:
                raise ValueError(f"Column '{c}' not in dataset.")
        a = self.df[col1]; b = self.df[col2]
        mask = ~(a.isna() | b.isna()); a=a[mask]; b=b[mask]
        a_num = pd.to_numeric(a, errors='coerce')
        b_num = pd.to_numeric(b, errors='coerce')
        if a_num.notna().all() and b_num.notna().all():
            return cohen_kappa_score(a_num, b_num, weights=weights)
        le = LabelEncoder()
        le.fit(pd.concat([a.astype(str), b.astype(str)], ignore_index=True))
        a_enc = le.transform(a.astype(str)); b_enc = le.transform(b.astype(str))
        return cohen_kappa_score(a_enc, b_enc, weights=weights)

    def scott_pi(self, col1: str, col2: str):
        if not col1 or not col2:
            raise ValueError("Select both columns for Scott's pi.")
        for c in (col1, col2):
            if c not in self.df.columns:
                raise ValueError(f"Column '{c}' not in dataset. Available: {list(self.df.columns)}")
        a = self.df[col1].astype(str); b = self.df[col2].astype(str)
        n = min(len(a), len(b)); a=a.iloc[:n]; b=b.iloc[:n]
        po = float((a.values == b.values).mean())
        pooled = pd.concat([a, b], ignore_index=True)
        p = pooled.value_counts(normalize=True)
        pe = float((p**2).sum())
        if pe == 1.0: return 1.0
        return (po - pe) / (1.0 - pe)

    def fleiss_kappa(self, cols: List[str]):
        raters = self.df[cols].dropna().astype(str)
        cats = sorted(pd.unique(raters.values.ravel()))
        cat_to_idx = {c:i for i,c in enumerate(cats)}
        table = np.zeros((len(raters), len(cats)), dtype=int)
        for i, (_, row) in enumerate(raters.iterrows()):
            codes = row.map(cat_to_idx).dropna().astype(int).values
            for code in codes:
                table[i, code] += 1
        kappa = float(sm_fleiss_kappa(table, method='fleiss'))
        return kappa, len(raters), len(cols), len(cats)

    def icc(self, cols: List[str]):
        if not HAS_PINGOUIN:
            raise RuntimeError("pingouin is not installed. Run: pip install pingouin")
        if cols is None or len(cols) < 2:
            raise ValueError("Provide at least 2 rater columns for ICC.")
        for c in cols:
            if c not in self.df.columns:
                raise ValueError(f"Column '{c}' not found in data.")
        dfw = self.df[cols].copy().dropna(how="any")
        if dfw.empty: raise ValueError("No complete rows across the selected rater columns.")
        dfw = dfw.reset_index().rename(columns={"index": "subject"})
        long = dfw.melt(id_vars="subject", value_vars=cols, var_name="rater", value_name="rating")
        out = pg.icc(data=long, targets="subject", raters="rater", ratings="rating")
        prefer = ["ICC1","ICC1k","ICC2","ICC2k","ICC3","ICC3k"]; out["order"] = out["Type"].apply(lambda t: prefer.index(t) if t in prefer else 999)
        return out.sort_values("order").drop(columns=["order"]).reset_index(drop=True)

    def kripp_alpha(self, cols: List[str], level: str = "nominal"):
        if not HAS_KRIPP:
            raise RuntimeError("krippendorff package is not installed. Run: pip install krippendorff")
        if cols is None or len(cols) < 2:
            raise ValueError("Provide at least 2 rater columns for Krippendorff’s alpha.")
        for c in cols:
            if c not in self.df.columns:
                raise ValueError(f"Column '{c}' not found in data.")
        mat = self.df[cols].T.values.tolist()
        alpha = kd.alpha(reliability_data=mat, level_of_measurement=level)
        return float(alpha)

    
    def survival(self, time: str, event: str, group: Optional[str], covars: List[str], report):
        if not HAS_LIFELINES: report.add_info("Survival unavailable","Install lifelines."); return
        from lifelines import KaplanMeierFitter, CoxPHFitter, WeibullAFTFitter
        df = self.df[[time, event] + ([group] if group else []) + covars].dropna()
        kmf = KaplanMeierFitter()
        if group and df[group].nunique()>1:
            fig = plt.figure()
            for g, sub in df.groupby(group):
                kmf.fit(sub[time], event_observed=sub[event], label=str(g)); kmf.plot_survival_function()
            plt.title("Kaplan–Meier"); report.add_figure(fig, "KM by group")
        else:
            kmf.fit(df[time], event_observed=df[event], label="All"); ax = kmf.plot_survival_function(); ax.set_title("Kaplan–Meier")
            report.add_figure(ax.get_figure(), "Kaplan–Meier")
        if covars:
            cph = CoxPHFitter(); cph.fit(df[[time,event]+covars], duration_col=time, event_col=event)
            report.add_table(cph.summary.reset_index(), "Cox PH")
            aft = WeibullAFTFitter(); aft.fit(df[[time,event]+covars], duration_col=time, event_col=event)
            report.add_table(aft.summary.reset_index(), "Weibull AFT")

    def diagnostic_curves(self, y: str, score: str, report):
        d = self.df[[y, score]].dropna()
        if d[y].nunique()!=2: report.add_info("Diagnostic error","Outcome must be binary."); return
        fpr, tpr, _ = roc_curve(d[y], d[score]); prec, rec, _ = precision_recall_curve(d[y], d[score])
        fig1=plt.figure(); plt.plot(fpr,tpr); plt.plot([0,1],[0,1],'--'); plt.xlabel("FPR"); plt.ylabel("TPR"); plt.title("ROC")
        fig2=plt.figure(); plt.plot(rec,prec); plt.xlabel("Recall"); plt.ylabel("Precision"); plt.title("PR curve")
        report.add_figure(fig1, "ROC"); report.add_figure(fig2, "PR"); brier = float(np.mean((d[score]-d[y])**2)); report.add_kv("Brier score", {"Brier": round(brier,4)})

    def arima_forecast(self, time_col: str, value_col: str, steps: int, report):
        from statsmodels.tsa.arima.model import ARIMA
        df = self.df[[time_col, value_col]].dropna().copy().sort_values(time_col)
        y = pd.Series(df[value_col].values, index=pd.to_datetime(df[time_col]))
        model = ARIMA(y, order=(1,0,1)).fit()
        fc = model.get_forecast(steps=steps); pred=fc.predicted_mean; ci=fc.conf_int()
        try:
            freq = pd.infer_freq(y.index); future_index = pd.date_range(y.index[-1], periods=steps+1, freq=freq or 'D')[1:]
        except Exception:
            future_index = pd.date_range(y.index[-1], periods=steps+1, freq='D')[1:]
        fig = plt.figure(); plt.plot(y.index, y.values, label="Observed"); plt.plot(future_index, pred.values, label="Forecast")
        plt.fill_between(future_index, ci.iloc[:,0].values, ci.iloc[:,1].values, alpha=0.2); plt.legend(); plt.title("ARIMA Forecast")
        report.add_figure(fig, "ARIMA")

    def ets_forecast(self, time_col: str, value_col: str, steps: int, report, seasonal=None):
        from statsmodels.tsa.holtwinters import ExponentialSmoothing
        df = self.df[[time_col, value_col]].dropna().copy().sort_values(time_col)
        y = pd.Series(df[value_col].values, index=pd.to_datetime(df[time_col]))
        model = ExponentialSmoothing(y, trend='add', seasonal=seasonal, seasonal_periods=(12 if seasonal else None)).fit()
        pred = model.forecast(steps)
        try:
            future_index = pd.date_range(y.index[-1], periods=steps+1, freq=pd.infer_freq(y.index) or 'D')[1:]
        except Exception:
            future_index = pd.date_range(y.index[-1], periods=steps+1, freq='D')[1:]
        fig = plt.figure(); plt.plot(y.index, y.values, label="Observed"); plt.plot(future_index, pred.values, label="ETS Forecast"); plt.legend()
        report.add_figure(fig, "ETS Forecast")

    def meta_dl(self, yi: str, vi: str, by: Optional[str], report):
        d = self.df[[yi,vi] + ([by] if by else [])].dropna().copy()
        def block(sub, label):
            y = sub[yi].values; v = sub[vi].values; w = 1/v
            yfix = np.sum(w*y)/np.sum(w); Q = np.sum(w*(y-yfix)**2); dfq=len(y)-1
            C = np.sum(w) - (np.sum(w**2)/np.sum(w)); tau2 = max(0.0, (Q-dfq)/C) if dfq>0 else 0.0
            wr = 1/(v+tau2); yr = np.sum(wr*y)/np.sum(wr); se = np.sqrt(1/np.sum(wr)); lo=yr-1.96*se; hi=yr+1.96*se
            I2 = max(0.0,(Q-dfq)/Q)*100 if Q>0 else 0.0
            return {"Group": label, "y_random": yr, "CI low": lo, "CI high": hi, "tau^2": tau2, "Q": Q, "df": dfq, "I^2%": I2}
        rows=[]
        if by:
            for g,sub in d.groupby(by): rows.append(block(sub, str(g)))
        rows.append(block(d, "Overall"))
        report.add_table(pd.DataFrame(rows).round(6), "Random-effects meta-analysis (DL)")

    def pca(self, cols: List[str], n_components: int, report):
        X = self.df[cols].dropna().astype(float).values; X = StandardScaler().fit_transform(X)
        p = PCA(n_components=n_components).fit(X)
        comp = pd.DataFrame(p.components_, columns=cols); expl = pd.DataFrame({"PC": np.arange(1, n_components+1), "ExplainedVar": p.explained_variance_ratio_})
        report.add_table(expl.round(4), "PCA explained variance"); report.add_table(comp.round(4), "Component loadings")
    def efa(self, cols: List[str], n_factors: int, report):
        if not HAS_FA: report.add_info("EFA unavailable","Install factor-analyzer."); return
        X = self.df[cols].dropna().astype(float).values; from factor_analyzer import FactorAnalyzer
        fa = FactorAnalyzer(n_factors=n_factors, rotation='varimax'); fa.fit(X)
        load = pd.DataFrame(fa.loadings_, columns=[f"F{i+1}" for i in range(n_factors)], index=cols); report.add_table(load.round(4), "EFA varimax loadings")
    def correspondence(self, a: str, b: str, report):
        if not HAS_PRINCE: report.add_info("Correspondence unavailable","Install 'prince'."); return
        tab = pd.crosstab(self.df[a], self.df[b]); import prince
        ca = prince.CA(n_components=2, n_iter=10, copy=True, check_input=True).fit(tab)
        r = ca.row_coordinates(tab); c = ca.column_coordinates(tab)
        report.add_table(r.round(4), "Row coordinates (CA)"); report.add_table(c.round(4), "Column coordinates (CA)")

    def lda(self, dv: str, predictors: List[str], report):
        d = self.df[[dv]+predictors].dropna(); y=d[dv].astype(str); X=d[predictors].astype(float)
        model = LDA().fit(X,y); acc = float(model.score(X,y)); report.add_kv("LDA training accuracy", {"Accuracy": round(acc,4)})
    def tree_cls(self, dv: str, predictors: List[str], report, criterion="gini"):
        from sklearn.tree import DecisionTreeClassifier
        d = self.df[[dv]+predictors].dropna(); y=d[dv].astype(str); X=d[predictors].astype(float)
        model = DecisionTreeClassifier(random_state=0, criterion=criterion, min_samples_leaf=5).fit(X,y); acc=float(model.score(X,y))
        report.add_kv(("CHAID-like tree" if criterion=='entropy' else "Decision tree")+" (training accuracy)", {"Accuracy": round(acc,4)})
    def rf_cls(self, dv: str, predictors: List[str], report):
        from sklearn.ensemble import RandomForestClassifier
        d = self.df[[dv]+predictors].dropna(); y=d[dv].astype(str); X=d[predictors].astype(float)
        model = RandomForestClassifier(n_estimators=200, random_state=0).fit(X,y); acc=float(model.score(X,y))
        report.add_kv("Random forest (training accuracy)", {"Accuracy": round(acc,4)})

    def kmeans(self, cols: List[str], k: int, report):
        X = self.df[cols].dropna().astype(float).values; X = StandardScaler().fit_transform(X)
        km = KMeans(n_clusters=k, n_init=10, random_state=0).fit(X); centers = pd.DataFrame(km.cluster_centers_, columns=cols)
        report.add_table(centers.round(4), "K-Means cluster centers")
    def agglomerative(self, cols: List[str], k: int, report):
        X = self.df[cols].dropna().astype(float).values; X = StandardScaler().fit_transform(X)
        agg = AgglomerativeClustering(n_clusters=k).fit(X); labels = pd.Series(agg.labels_, name="cluster")
        report.add_table(pd.DataFrame(labels.value_counts()), "Agglomerative cluster sizes")
    def auto_kmeans(self, cols: List[str], report):
        X = self.df[cols].dropna().astype(float).values; Xs = StandardScaler().fit_transform(X)
        best_k, best_s = None, -1
        for k in range(2, 11):
            km = KMeans(n_clusters=k, n_init=10, random_state=0).fit(Xs)
            try:
                s = silhouette_score(Xs, km.labels_)
                if s > best_s: best_s, best_k = s, k
            except Exception:
                pass
        report.add_kv("Auto-KMeans selection", {"k": best_k, "silhouette": round(float(best_s),4)})


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle(WINDOW_TITLE); self.resize(1420, 980)
        self.df=None; self.model=PandasModel(pd.DataFrame())
        self.tabs = QTabWidget(); self.setCentralWidget(self.tabs)
        self._build_welcome_tab(); self._build_data_tab(); self._build_analyze_tab(); self._build_results_tab(); self._build_help_tab(); self.tabs.setCurrentIndex(0)
        self.report = ReportBuilder(); self.statusBar().showMessage("Ready. Load CSV/XLSX to begin."); self.setAcceptDrops(True)

    
    
    def _build_welcome_tab(self):
        tab = QWidget(); layout = QVBoxLayout(tab)

        
        self.readme_view = QTextBrowser()
        self.readme_view.setOpenExternalLinks(True)
        readme_html = None
        
        rd_candidates = [os.path.join(_app_dir(), "README.md"), os.path.join(os.getcwd(), "README.md")]
        for p in rd_candidates:
            if os.path.exists(p):
                try:
                    with open(p, "r", encoding="utf-8") as f:
                        txt = f.read()
                    readme_html = "<h2>README</h2><pre style='white-space:pre-wrap;'>" + html.escape(txt) + "</pre>"
                    break
                except Exception: pass
        if not readme_html:
            readme_html = "<h2>README</h2><p>README.md not found next to the application.</p>"
        self.readme_view.setHtml(readme_html)

        
        self.reqs_view = QTextBrowser()
        reqs_html = ["<h2>Requirements (status on this system)</h2><table><tr><th>Package</th><th>Required</th><th>Installed</th></tr>"]
        req_file = None
        rq_candidates = [os.path.join(_app_dir(), "requirements.txt"), os.path.join(os.getcwd(), "requirements.txt")]
        for p in rq_candidates:
            if os.path.exists(p):
                req_file = p; break
        pkgs = []
        if req_file:
            with open(req_file, "r", encoding="utf-8") as f:
                for line in f:
                    line=line.strip()
                    if not line or line.startswith("#"): continue
                    req = line
                    
                    name = re.split(r"[<>=!]", line, maxsplit=1)[0].strip()
                    if not name: name = line.strip()
                    try:
                        ver = ilm.version(name)
                        status = ver
                    except Exception:
                        status = "<span style='color:#b91c1c;'>Missing</span>"
                        ver = "—"
                    reqs_html.append(f"<tr><td>{html.escape(name)}</td><td>{html.escape(req)}</td><td>{status}</td></tr>")
                    pkgs.append(name)
        else:
            reqs_html.append("<tr><td colspan='3'>requirements.txt not found.</td></tr>")
        reqs_html.append("</table>")
        self.reqs_view.setHtml("".join(reqs_html))

        
        self.samples_list = QListWidget()
        samp_dir = None
        for p in [os.path.join(_app_dir(), "samples"), os.path.join(os.getcwd(), "samples")]:
            if os.path.isdir(p):
                samp_dir = p; break
        self._samples_dir = samp_dir
        if samp_dir:
            for fn in sorted(os.listdir(samp_dir)):
                if fn.lower().endswith((".csv",".xlsx",".xls")):
                    item = QListWidgetItem(fn)
                    self.samples_list.addItem(item)

        btns = QHBoxLayout()
        self.btn_load_sample = QPushButton("Open selected sample → Data tab")
        self.btn_load_sample.clicked.connect(self._open_selected_sample)
        btns.addWidget(self.btn_load_sample); btns.addStretch(1)

        
        layout.addWidget(QLabel("<b>Welcome to Statilytics Studio v1.0</b>"))
        layout.addWidget(self.readme_view)
        layout.addWidget(self.reqs_view)
        layout.addWidget(QLabel("<b>Sample datasets</b> (double-click or select and click Open):"))
        layout.addWidget(self.samples_list)
        layout.addLayout(btns)

        
        self.samples_list.itemDoubleClicked.connect(self._open_selected_sample)

        self.tabs.addTab(tab, "Welcome / Resources")

    def _open_selected_sample(self):
        item = self.samples_list.currentItem()
        if not item or not self._samples_dir:
            QMessageBox.information(self, "Samples", "No sample selected.")
            return
        path = os.path.join(self._samples_dir, item.text())
        try:
            if path.lower().endswith((".xlsx",".xls")):
                book = pd.read_excel(path, sheet_name=None)
                df = book.get("Data", next(iter(book.values())))
            else:
                df = pd.read_csv(path)
            self.set_dataframe(df)
            self.tabs.setCurrentIndex(1)  # switch to Data tab
            QMessageBox.information(self, "Sample loaded", f"Loaded: {os.path.basename(path)}")
        except Exception as e:
            QMessageBox.critical(self, "Sample error", str(e))

    def _build_data_tab(self):
        tab = QWidget(); lay = QVBoxLayout(tab)
        self.table = QTableView(); self.table.setModel(self.model)
        self.table.setHorizontalScrollMode(QAbstractItemView.ScrollPerPixel)
        self.table.setVerticalScrollMode(QAbstractItemView.ScrollPerPixel)
        self.table.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.table.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.table.setWordWrap(False)
        top = QHBoxLayout()
        self.btn_open = QPushButton("Open CSV/XLSX…"); self.btn_open.clicked.connect(self.open_file)
        self.lbl_shape = QLabel("No data loaded")
        top.addWidget(self.btn_open); top.addStretch(1); top.addWidget(self.lbl_shape)
        lay.addLayout(top); lay.addWidget(self.table)
        self.tabs.addTab(tab, "Data")

    def _build_analyze_tab(self):
        tab = QWidget(); layout = QHBoxLayout(tab)

        controls_group = QGroupBox("Procedures & Options")
        form = QFormLayout(controls_group)
        form.setLabelAlignment(Qt.AlignLeft)
        form.setFieldGrowthPolicy(QFormLayout.AllNonFixedFieldsGrow)

        self.cmb_proc = QComboBox(); self.cmb_proc.addItems([
            "Describe data","Crosstab (chi-square/Fisher)",
            "t-test (one-sample)","t-test (independent)","t-test (paired)",
            "ANOVA (one-way)","Repeated-measures ANOVA","ANCOVA","Tukey HSD",
            "Mann–Whitney U","Wilcoxon signed-rank","Kruskal–Wallis","Friedman",
            "Correlation (Pearson)","Correlation (Spearman)","Correlation (Kendall)",
            "OLS regression","Logistic regression","Multinomial logit","Ordinal regression (logit)",
            "GLM (Poisson)","GLM (Negative binomial)","GLM (Gamma)",
            "Linear mixed model",
            "Cronbach alpha","Cohen kappa","Weighted kappa (linear)","Weighted kappa (quadratic)",
            "Fleiss' kappa (multi-rater)","ICC",
            "Krippendorff’s alpha (nominal)","Krippendorff’s alpha (ordinal)","Krippendorff’s alpha (interval)",
            "Survival (KM + Cox + Weibull AFT)",
            "ROC / PR / Brier",
            "ARIMA forecast","ETS forecast",
            "Meta-analysis (DL random-effects)",
            "IPTW (ATE)","Difference-in-Differences",
            "PCA","EFA (varimax)","Correspondence analysis",
            "LDA (classification)","Decision tree (classification)","CHAID-like tree","Random forest (classification)",
            "K-Means","Agglomerative clustering","Auto-KMeans (TwoStep-inspired)",
            "MANOVA / MANCOVA"
        ]); form.addRow("Procedure", self.cmb_proc)

        self.cmb_y = QComboBox(); form.addRow("Y / Outcome", self.cmb_y)
        self.cmb_x = QComboBox(); form.addRow("X / Second Var", self.cmb_x)
        self.cmb_group = QComboBox(); form.addRow("Group / Factor", self.cmb_group)
        self.cmb_group2 = QComboBox(); form.addRow("Second Group (optional)", self.cmb_group2)
        self.txt_predictors = QLineEdit(); form.addRow("Predictors (comma/semicolon/space, quotes OK)", self.txt_predictors)
        self.cmb_time = QComboBox(); form.addRow("Time (survival/TS)", self.cmb_time)
        self.cmb_event = QComboBox(); form.addRow("Event (0/1)", self.cmb_event)
        self.cmb_cluster = QComboBox(); form.addRow("Cluster (LMM)", self.cmb_cluster)
        self.cmb_subject = QComboBox(); form.addRow("Subject ID (RM ANOVA)", self.cmb_subject)
        self.txt_within = QLineEdit(); form.addRow("Within-factor name (RM ANOVA)", self.txt_within)
        self.txt_multi_y = QLineEdit(); form.addRow("Y's (comma for MANOVA)", self.txt_multi_y)
        self.txt_const = QLineEdit(); form.addRow("Constant / μ", self.txt_const)
        self.txt_levels = QLineEdit(); form.addRow("Group levels", self.txt_levels)
        self.spn_steps = QSpinBox(); self.spn_steps.setRange(1, 500); self.spn_steps.setValue(12); form.addRow("Forecast steps", self.spn_steps)
        self.spn_k = QSpinBox(); self.spn_k.setRange(2, 20); self.spn_k.setValue(3); form.addRow("Components / Clusters (k)", self.spn_k)
        self.chk_fisher = QCheckBox("Use Fisher exact (2x2)"); form.addRow(self.chk_fisher)
        self.btn_run = QPushButton("Run"); self.btn_run.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed); self.btn_run.clicked.connect(self.run); form.addRow(self.btn_run)

        left_container = QWidget(); left_v = QVBoxLayout(left_container); left_v.addWidget(controls_group); left_v.addStretch(1)
        scroll_left = QScrollArea(); scroll_left.setWidgetResizable(True)
        scroll_left.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        scroll_left.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        scroll_left.setWidget(left_container)
        scroll_left.setMinimumWidth(430)

        right = QWidget(); rlay = QVBoxLayout(right)
        self.txt_report = QTextEdit(); self.txt_report.setReadOnly(True)
        self.txt_report.setLineWrapMode(QTextEdit.NoWrap)
        self.txt_report.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.txt_report.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        rlay.addWidget(self.txt_report)

        layout.addWidget(scroll_left, 0)
        layout.addWidget(right, 1)
        self.tabs.addTab(tab, "Analyze")

    
    def _build_help_tab(self):
        tab = QWidget(); layout = QVBoxLayout(tab)
        self.help_view = QTextBrowser()
        self.help_view.setOpenExternalLinks(True)

        
        app_dir = None
        try:
            app_dir = os.path.dirname(os.path.abspath(__file__))
        except Exception:
            app_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
        if getattr(sys, "frozen", False):
            app_dir = getattr(sys, "_MEIPASS", app_dir)

        candidates = [
            os.path.join(app_dir, "USER_MANUAL.html"),
            os.path.join(os.getcwd(), "USER_MANUAL.html")
        ]
        manual_path = next((p for p in candidates if os.path.exists(p)), None)

        if manual_path:
            self.help_view.setSource(QtCore.QUrl.fromLocalFile(manual_path))
        else:
            self.help_view.setHtml(f"<h2>{APP_NAME}</h2><p>{COPYRIGHT}</p><p><b>User Manual not found.</b> Place <code>USER_MANUAL.html</code> in the same folder as the application.</p>")

        layout.addWidget(self.help_view)
        self.tabs.addTab(tab, "Help / Manual")

    def _build_results_tab(self):
        tab = QWidget(); layout = QVBoxLayout(tab)
        top = QHBoxLayout()
        self.btn_export_html = QPushButton("Export HTML…"); self.btn_export_html.clicked.connect(self.export_html)
        self.btn_export_docx = QPushButton("Export DOCX…"); self.btn_export_docx.clicked.connect(self.export_docx)
        top.addWidget(self.btn_export_html); top.addWidget(self.btn_export_docx); top.addStretch(1)
        self.txt_final = QTextEdit(); self.txt_final.setReadOnly(True); self.txt_final.setLineWrapMode(QTextEdit.NoWrap)
        self.txt_final.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.txt_final.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        layout.addLayout(top); layout.addWidget(self.txt_final)
        self.tabs.addTab(tab, "Results & Export")

    
    def open_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "Open Data", os.getcwd(), "Data (*.csv *.xlsx *.xls)")
        if not path: return
        try:
            if path.lower().endswith((".xlsx",".xls")):
                book = pd.read_excel(path, sheet_name=None)
                df = book.get("Data", next(iter(book.values())))
            else:
                df = pd.read_csv(path)
            self.set_dataframe(df)
        except Exception as e:
            QMessageBox.critical(self,"File error",str(e))

    def set_dataframe(self, df: pd.DataFrame):
        self.df=df; self.model.setDataFrame(df); self.lbl_shape.setText(f"Rows: {df.shape[0]} | Cols: {df.shape[1]}")
        cols = ["—"] + list(df.columns)
        combos = [self.cmb_y, self.cmb_x, self.cmb_group, self.cmb_group2, self.cmb_time, self.cmb_event, self.cmb_cluster, self.cmb_subject]
        for cmb in combos:
            cmb.clear(); cmb.addItems(cols)
        self.report = ReportBuilder(); self.report.add_info("Data loaded", f"Shape: {df.shape}")
        self._refresh()

    
    def _split_simple(self, txt: str):
        return [t.strip() for t in re.split(r"[\\s,;]+", txt or "") if t.strip()]

    def _parse_cols(self, txt: str):
        """Parse Predictors allowing commas/semicolons/spaces and quoted names."""
        if not txt or not txt.strip():
            return []
        s = txt.replace(";", ",").strip()
        cols = []
        try:
            cols = next(csv.reader([s], skipinitialspace=True))
        except Exception:
            cols = []
        if len(cols) <= 1 and ("," not in s):
            cols = [t for t in re.split(r"\\s+", s) if t]
        return [c.strip().strip('"').strip("'") for c in cols if c.strip().strip('"').strip("'")]

    def _val(self, combo: QComboBox):
        return None if combo.currentText()=="—" else combo.currentText()

    
    def run(self):
        if self.df is None: QMessageBox.warning(self,"No data","Load a dataset first."); return
        E = Engine(self.df); R = ReportBuilder()
        p = self.cmb_proc.currentText()
        val = self._val

        try:
            if p=="Describe data":
                R.add_table(E.describe().round(4), "Descriptives")

            elif p=="Crosstab (chi-square/Fisher)":
                a = val(self.cmb_y); b = val(self.cmb_x); 
                if not a or not b: raise ValueError("Pick both Y and X for crosstab.")
                tab, info = E.crosstab(a,b,fisher=self.chk_fisher.isChecked()); R.add_table(tab,"Crosstab"); R.add_kv("Association test", info)

            elif p=="t-test (one-sample)":
                if not val(self.cmb_y): raise ValueError("Select Y for one-sample t.")
                R.add_kv("One-sample t", E.t_one_sample(val(self.cmb_y), float(self.txt_const.text() or 0.0)))

            elif p=="t-test (independent)":
                if not val(self.cmb_y) or not val(self.cmb_group): raise ValueError("Pick Y and Group for independent t.")
                levels = (self.txt_levels.text() or "").split(","); 
                if len(levels)<2: raise ValueError("Provide two group levels, e.g., A,B")
                R.add_kv("Independent t", E.t_ind(val(self.cmb_y), val(self.cmb_group), levels[0].strip(), levels[1].strip(), equal_var=False))

            elif p=="t-test (paired)":
                if not val(self.cmb_y) or not val(self.cmb_x): raise ValueError("Pick Y and X for paired t.")
                R.add_kv("Paired t", E.t_paired(val(self.cmb_y), val(self.cmb_x)))

            elif p=="ANOVA (one-way)":
                if not val(self.cmb_y) or not val(self.cmb_group): raise ValueError("Pick Y and Group for ANOVA.")
                F,pv = E.anova_oneway(val(self.cmb_y), val(self.cmb_group)); R.add_kv("One-way ANOVA", {"F": round(F,4),"p": round(pv,6)})

            elif p=="Repeated-measures ANOVA":
                if not val(self.cmb_y) or not val(self.cmb_subject) or not self.txt_within.text().strip(): raise ValueError("Y, Subject, and Within name required.")
                R.add_table(E.anova_rm(val(self.cmb_y), val(self.cmb_subject), self.txt_within.text().strip()), "Repeated-measures ANOVA")

            elif p=="ANCOVA":
                if not val(self.cmb_y) or not val(self.cmb_group): raise ValueError("Pick Y and Group for ANCOVA.")
                m = E.ancova(val(self.cmb_y), val(self.cmb_group), self._split_simple(self.txt_predictors.text())); R.add_table(m.summary2().tables[1].reset_index(), "ANCOVA (OLS table)")

            elif p=="Tukey HSD":
                if not val(self.cmb_y) or not val(self.cmb_group): raise ValueError("Pick Y and Group for Tukey.")
                R.add_table(E.tukey_hsd(val(self.cmb_y), val(self.cmb_group)), "Tukey HSD pairwise")

            elif p=="Mann–Whitney U":
                if not val(self.cmb_y) or not val(self.cmb_group): raise ValueError("Pick Y and Group for Mann–Whitney.")
                levels=(self.txt_levels.text() or "").split(","); 
                if len(levels)<2: raise ValueError("Provide two group levels, e.g., A,B")
                R.add_kv("Mann–Whitney U", E.mannwhitney(val(self.cmb_y), val(self.cmb_group), levels[0].strip(), levels[1].strip()))

            elif p=="Wilcoxon signed-rank":
                if not val(self.cmb_y) or not val(self.cmb_x): raise ValueError("Pick Y and X for Wilcoxon.")
                R.add_kv("Wilcoxon signed-rank", E.wilcoxon(val(self.cmb_y), val(self.cmb_x)))

            elif p=="Kruskal–Wallis":
                if not val(self.cmb_y) or not val(self.cmb_group): raise ValueError("Pick Y and Group for Kruskal–Wallis.")
                R.add_kv("Kruskal–Wallis", E.kruskal(val(self.cmb_y), val(self.cmb_group)))

            elif p=="Friedman":
                cols = self._split_simple(self.txt_predictors.text())
                if len(cols)<3: raise ValueError("Provide 3 columns for Friedman.")
                R.add_kv("Friedman", E.friedman(cols))

            elif p=="Correlation (Pearson)":
                cols = self._split_simple(self.txt_predictors.text()); 
                if len(cols)<2: raise ValueError("Provide at least 2 columns for correlation.")
                R.add_table(E.correlations(cols, 'pearson').round(4),"Pearson correlation")

            elif p=="Correlation (Spearman)":
                cols = self._split_simple(self.txt_predictors.text()); 
                if len(cols)<2: raise ValueError("Provide at least 2 columns for correlation.")
                R.add_table(E.correlations(cols, 'spearman').round(4),"Spearman correlation")

            elif p=="Correlation (Kendall)":
                cols = self._split_simple(self.txt_predictors.text()); 
                if len(cols)<2: raise ValueError("Provide at least 2 columns for correlation.")
                R.add_table(E.correlations(cols, 'kendall').round(4),"Kendall correlation")

            elif p=="OLS regression":
                if not val(self.cmb_y): raise ValueError("Pick Y for OLS.")
                m=E.ols(val(self.cmb_y), self._split_simple(self.txt_predictors.text())); R.add_table(m.summary2().tables[1].reset_index(), "OLS coefficients")

            elif p=="Logistic regression":
                if not val(self.cmb_y): raise ValueError("Pick Y for logistic.")
                m=E.logit(val(self.cmb_y), self._split_simple(self.txt_predictors.text())); tab = m.summary2().tables[1].reset_index(); tab["OR"]=np.exp(tab["Coef."]); R.add_table(tab, "Logistic coefficients (OR)")

            elif p=="Multinomial logit":
                if not val(self.cmb_y): raise ValueError("Pick Y for multinomial.")
                m=E.mlogit(val(self.cmb_y), self._split_simple(self.txt_predictors.text()))
                for i,t in enumerate(m.summary().tables):
                    if hasattr(t,'data'):
                        import pandas as pd; R.add_table(pd.DataFrame(t.data[1:], columns=t.data[0]), f"MNLogit table {i+1}")

            elif p=="Ordinal regression (logit)":
                if not val(self.cmb_y): raise ValueError("Pick Y for ordinal regression.")
                m=E.ordinal(val(self.cmb_y), self._split_simple(self.txt_predictors.text())); R.add_table(m.summary().tables[1], "Ordered logit coefficients")

            elif p=="GLM (Poisson)":
                if not val(self.cmb_y): raise ValueError("Pick Y for GLM.")
                m=E.glm(val(self.cmb_y), self._split_simple(self.txt_predictors.text()), 'poisson'); R.add_table(m.summary2().tables[1].reset_index(), "Poisson GLM")

            elif p=="GLM (Negative binomial)":
                if not val(self.cmb_y): raise ValueError("Pick Y for GLM.")
                m=E.glm(val(self.cmb_y), self._split_simple(self.txt_predictors.text()), 'negbin'); R.add_table(m.summary2().tables[1].reset_index(), "Negative binomial GLM")

            elif p=="GLM (Gamma)":
                if not val(self.cmb_y): raise ValueError("Pick Y for GLM.")
                m=E.glm(val(self.cmb_y), self._split_simple(self.txt_predictors.text()), 'gamma'); R.add_table(m.summary2().tables[1].reset_index(), "Gamma GLM (log link)")

            elif p=="Linear mixed model":
                if not val(self.cmb_y) or not val(self.cmb_cluster): raise ValueError("Pick Y and Cluster for LMM.")
                m=E.lmm(val(self.cmb_y), self._split_simple(self.txt_predictors.text()), val(self.cmb_cluster)); R.add_table(m.summary().tables[1], "LMM fixed effects")

            elif p=="Cronbach alpha":
                cols = self._split_simple(self.txt_predictors.text())
                if len(cols)<2: raise ValueError("Provide at least 2 item columns for alpha.")
                R.add_kv("Cronbach alpha", {"alpha": round(float(E.cronbach_alpha(cols)),4)})

            elif p=="Cohen kappa":
                y = val(self.cmb_y); x = val(self.cmb_x)
                if not y or not x: raise ValueError("Select both Y and X columns for Cohen's kappa.")
                R.add_kv("Cohen kappa", {"kappa": round(float(E.kappa(y, x)),4)})

            elif p=="Weighted kappa (linear)":
                y = val(self.cmb_y); x = val(self.cmb_x)
                if not y or not x: raise ValueError("Select both Y and X columns for weighted kappa.")
                R.add_kv("Weighted Cohen kappa (linear)", {"kappa_w_linear": round(float(E.weighted_kappa(y, x, weights='linear')),4)})

            elif p=="Weighted kappa (quadratic)":
                y = val(self.cmb_y); x = val(self.cmb_x)
                if not y or not x: raise ValueError("Select both Y and X columns for weighted kappa.")
                R.add_kv("Weighted Cohen kappa (quadratic)", {"kappa_w_quadratic": round(float(E.weighted_kappa(y, x, weights='quadratic')),4)})

            elif p=="Fleiss' kappa (multi-rater)":
                cols = self._parse_cols(self.txt_predictors.text())
                if len(cols) < 2:
                    y = val(self.cmb_y); x = val(self.cmb_x)
                    if y and x:
                        k = E.scott_pi(y, x)
                        R.add_kv("Scott's pi (auto: used Y and X)", {"pi": round(float(k), 4), "rater1": y, "rater2": x})
                    else:
                        R.add_info("Fleiss' kappa", "Provide at least 2 rater columns in Predictors OR set Y and X (two raters). With exactly 2, result = Scott's π.")
                elif len(cols) == 2:
                    k = E.scott_pi(cols[0], cols[1])
                    R.add_kv("Scott's pi (Fleiss with 2 raters)", {"pi": round(float(k), 4), "rater1": cols[0], "rater2": cols[1]})
                else:
                    kappa, n_subj, n_raters, n_cats = E.fleiss_kappa(cols)
                    R.add_kv("Fleiss' kappa (multi-rater)", {
                        "kappa": round(kappa, 4),
                        "subjects (rows used)": n_subj,
                        "raters (columns)": n_raters,
                        "categories detected": n_cats
                    })

            elif p=="ICC":
                cols = self._parse_cols(self.txt_predictors.text())
                if len(cols) < 2:
                    raise ValueError("For ICC, list 2+ rater columns in Predictors (e.g., r1,r2,r3).")
                R.add_table(E.icc(cols).round(6), "Intraclass correlation")

            elif p=="Krippendorff’s alpha (nominal)":
                cols = self._parse_cols(self.txt_predictors.text())
                if len(cols) < 2:
                    y = val(self.cmb_y); x = val(self.cmb_x)
                    if y and x:
                        cols = [y, x]
                    else:
                        raise ValueError("Provide 2+ rater columns in Predictors or set Y and X (two raters).")
                a = E.kripp_alpha(cols, level="nominal"); R.add_kv("Krippendorff’s α (nominal)", {"alpha": round(float(a),4), "raters": ", ".join(cols)})

            elif p=="Krippendorff’s alpha (ordinal)":
                cols = self._parse_cols(self.txt_predictors.text())
                if len(cols) < 2:
                    y = val(self.cmb_y); x = val(self.cmb_x)
                    if y and x:
                        cols = [y, x]
                    else:
                        raise ValueError("Provide 2+ rater columns in Predictors or set Y and X (two raters).")
                a = E.kripp_alpha(cols, level="ordinal"); R.add_kv("Krippendorff’s α (ordinal)", {"alpha": round(float(a),4), "raters": ", ".join(cols)})

            elif p=="Krippendorff’s alpha (interval)":
                cols = self._parse_cols(self.txt_predictors.text())
                if len(cols) < 2:
                    y = val(self.cmb_y); x = val(self.cmb_x)
                    if y and x:
                        cols = [y, x]
                    else:
                        raise ValueError("Provide 2+ rater columns in Predictors or set Y and X (two raters).")
                a = E.kripp_alpha(cols, level="interval"); R.add_kv("Krippendorff’s α (interval)", {"alpha": round(float(a),4), "raters": ", ".join(cols)})

            elif p=="Survival (KM + Cox + Weibull AFT)":
                if not val(self.cmb_time) or not val(self.cmb_event): raise ValueError("Pick Time and Event for survival.")
                E.survival(val(self.cmb_time), val(self.cmb_event), val(self.cmb_group), self._split_simple(self.txt_predictors.text()), R)

            elif p=="ROC / PR / Brier":
                if not val(self.cmb_y) or not val(self.cmb_x): raise ValueError("Pick Y (binary true) and X (score) for ROC/PR.")
                E.diagnostic_curves(val(self.cmb_y), val(self.cmb_x), R)

            elif p=="ARIMA forecast":
                if not val(self.cmb_time) or not (val(self.cmb_y) or val(self.cmb_x)): raise ValueError("Pick Time and Y/X for ARIMA.")
                E.arima_forecast(val(self.cmb_time), val(self.cmb_y) or val(self.cmb_x), self.spn_steps.value(), R)

            elif p=="ETS forecast":
                if not val(self.cmb_time) or not (val(self.cmb_y) or val(self.cmb_x)): raise ValueError("Pick Time and Y/X for ETS.")
                E.ets_forecast(val(self.cmb_time), val(self.cmb_y) or val(self.cmb_x), self.spn_steps.value(), R, seasonal=None)

            elif p=="Meta-analysis (DL random-effects)":
                if not val(self.cmb_y) or not val(self.cmb_x): raise ValueError("Provide yi (effect) in Y and vi (variance) in X.")
                E.meta_dl(val(self.cmb_y), val(self.cmb_x), val(self.cmb_group), R)

            elif p=="IPTW (ATE)":
                y=val(self.cmb_y); a=val(self.cmb_group); cov=self._split_simple(self.txt_predictors.text())
                if not y or not a or not cov: raise ValueError("Y, treatment Group, and covariates are required for IPTW.")
                d = self.df[[y, a]+cov].dropna().copy(); X=d[cov].values; T=d[a].astype(int).values; Y=d[y].astype(int).values
                lr = LogisticRegression(max_iter=200).fit(X,T); ps = np.clip(lr.predict_proba(X)[:,1], 1e-3, 1-1e-3)
                w = T/ps + (1-T)/(1-ps); mdl = sm.GLM(Y, sm.add_constant(T), family=sm.families.Binomial(), freq_weights=w).fit()
                OR=np.exp(mdl.params[1]); lo,hi=np.exp(mdl.conf_int().loc[1]); R.add_kv("IPTW (ATE) weighted logit", {"OR": round(float(OR),4), "CI low": round(float(lo),4), "CI high": round(float(hi),4)})

            elif p=="Difference-in-Differences":
                y=val(self.cmb_y); tr=val(self.cmb_group); tm=val(self.cmb_group2); cov=self._split_simple(self.txt_predictors.text())
                if not y or not tr or not tm: raise ValueError("Y, treatment, and time group columns are required for DiD.")
                formula=f"{y} ~ {tr} + {tm} + {tr}:{tm}" + ((" + " + " + ".join(cov)) if cov else "")
                m=smf.ols(formula, data=self.df.dropna(subset=[y,tr,tm]+cov)).fit(); R.add_table(m.summary2().tables[1].reset_index(), "DiD coefficients")

            elif p=="PCA":
                cols = self._split_simple(self.txt_predictors.text()); 
                if len(cols)<2: raise ValueError("Provide at least 2 numeric columns for PCA.")
                E.pca(cols, self.spn_k.value(), R)

            elif p=="EFA (varimax)":
                cols = self._split_simple(self.txt_predictors.text()); 
                if len(cols)<2: raise ValueError("Provide at least 2 numeric columns for EFA.")
                E.efa(cols, self.spn_k.value(), R)

            elif p=="Correspondence analysis":
                a = val(self.cmb_group); b = val(self.cmb_group2)
                if not a or not b: raise ValueError("Pick two categorical columns for correspondence analysis.")
                E.correspondence(a, b, R)

            elif p=="LDA (classification)":
                if not val(self.cmb_y): raise ValueError("Pick Y for LDA.")
                E.lda(val(self.cmb_y), self._split_simple(self.txt_predictors.text()), R)

            elif p=="Decision tree (classification)":
                if not val(self.cmb_y): raise ValueError("Pick Y for decision tree.")
                E.tree_cls(val(self.cmb_y), self._split_simple(self.txt_predictors.text()), R, criterion="gini")

            elif p=="CHAID-like tree":
                if not val(self.cmb_y): raise ValueError("Pick Y for CHAID-like tree.")
                E.tree_cls(val(self.cmb_y), self._split_simple(self.txt_predictors.text()), R, criterion="entropy")

            elif p=="Random forest (classification)":
                if not val(self.cmb_y): raise ValueError("Pick Y for random forest.")
                E.rf_cls(val(self.cmb_y), self._split_simple(self.txt_predictors.text()), R)

            elif p=="K-Means":
                cols = self._split_simple(self.txt_predictors.text()); 
                if len(cols)<2: raise ValueError("Provide numeric columns for K-Means.")
                E.kmeans(cols, self.spn_k.value(), R)

            elif p=="Agglomerative clustering":
                cols = self._split_simple(self.txt_predictors.text()); 
                if len(cols)<2: raise ValueError("Provide numeric columns for Agglomerative clustering.")
                E.agglomerative(cols, self.spn_k.value(), R)

            elif p=="Auto-KMeans (TwoStep-inspired)":
                cols = self._split_simple(self.txt_predictors.text()); 
                if len(cols)<2: raise ValueError("Provide numeric columns for Auto-KMeans.")
                E.auto_kmeans(cols, R)

            elif p=="MANOVA / MANCOVA":
                ys = [t.strip() for t in (self.txt_multi_y.text() or "").split(",") if t.strip()]
                res = E.manova(ys, self._split_simple(self.txt_predictors.text())); R.add_info("MANOVA / MANCOVA", f"<pre>{res}</pre>")

            else:
                raise ValueError("Unknown procedure.")
        except Exception as e:
            R.add_info("Run error", f"{e}<br><pre>{traceback.format_exc()}</pre>")

        self.report = R; self._refresh()

    
    def _refresh(self):
        html = self.report.html(); 
        self.txt_report.setHtml(html)
        if hasattr(self, "txt_final"):
            self.txt_final.setHtml(html)

    def export_html(self):
        p,_ = QFileDialog.getSaveFileName(self,"Save HTML", os.path.join(os.getcwd(),"report.html"), "HTML (*.html)")
        if not p: return
        with open(p,'w',encoding='utf-8') as f: f.write(self.report.html())
        QMessageBox.information(self,"Saved",f"Saved to {p}")

    def export_docx(self):
        try:
            from docx import Document
        except Exception:
            QMessageBox.information(self,"DOCX unavailable","Install python-docx to export DOCX."); return
        p,_ = QFileDialog.getSaveFileName(self,"Save DOCX", os.path.join(os.getcwd(),"report.docx"), "DOCX (*.docx)")
        if not p: return
        doc = Document(); doc.add_heading(APP_NAME, level=1); doc.add_paragraph(COPYRIGHT)
        for para in self.txt_report.toPlainText().split("\n\n"): doc.add_paragraph(para)
        doc.save(p); QMessageBox.information(self,"Saved",f"Saved to {p}")


def main():
    app = QApplication(sys.argv); app.setApplicationName(APP_NAME); app.setStyle("Fusion")
    pal = app.palette(); pal.setColor(QtGui.QPalette.Window, QtGui.QColor(248,250,252)); pal.setColor(QtGui.QPalette.Base, Qt.white); app.setPalette(pal)
    w = MainWindow(); w.show(); sys.exit(app.exec())

if __name__ == "__main__":
    main()
